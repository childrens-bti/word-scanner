import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
import fitz  # PyMuPDF
import io
import re
import pandas as pd
import os

st.set_page_config(page_title="Word Scanner", layout="wide")
st.title("🔍 Word Scanner for DOCX and PDF")
st.markdown("Upload a Word or PDF document to search for words within the document.")
st.markdown("""
**Sources of Banned Word Lists:**
- [Thoughtcrime Checker](https://thoughtcrime-checker.com/banned-words.html)
- [PEN America Banned Words List](https://pen.org/banned-words-list/)
- [Grant Writing & Funding: Banned and Trigger Words](https://grantwritingandfunding.com/banned-and-trigger-words-in-federal-grant-writing-in-the-trump-administration-2-0/)
- [U.S. Senate Report of NSF Banned Words](https://www.commerce.senate.gov/services/files/4BD2D522-2092-4246-91A5-58EEF99750BC)
""")

# --- Helper Functions ---
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file):
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text() for page in pdf])

def highlight_banned_words(text, banned_words):
    pattern = r"\b(" + "|".join(map(re.escape, banned_words)) + r")\b"
    matches = list(re.finditer(pattern, text, flags=re.IGNORECASE))

    spans = []
    for m in matches:
        span_start = max(0, m.start() - 40)
        span_end = min(len(text), m.end() + 40)
        context = text[span_start:span_end].replace(m.group(), f"**{m.group()}**")
        spans.append({
            "word": m.group(),
            "start_pos": m.start(),
            "end_pos": m.end(),
            "context": context
        })
    return spans

def generate_word_doc(results, full_text, banned_words):
    doc = Document()
    doc.add_heading('Full Document with Highlighted Banned Words', level=1)

    # Highlight full text
    p = doc.add_paragraph()
    pattern = re.compile(r"\b(" + "|".join(map(re.escape, banned_words)) + r")\b", flags=re.IGNORECASE)
    last_end = 0
    for match in pattern.finditer(full_text):
        # Add text before the match
        if last_end < match.start():
            p.add_run(full_text[last_end:match.start()])
        # Add highlighted match
        highlighted = p.add_run(full_text[match.start():match.end()])
        highlighted.font.highlight_color = 3  # Yellow highlight
        last_end = match.end()
    # Add any remaining text
    p.add_run(full_text[last_end:])

    # Add a page break and summary context
    doc.add_page_break()
    doc.add_heading('Context Summary of Banned Words', level=2)
    for entry in results:
        p = doc.add_paragraph()
        p.add_run("Word: ").bold = True
        p.add_run(entry['word'] + "\n")
        p.add_run("Location: ").bold = True
        p.add_run(f"{entry['start_pos']} - {entry['end_pos']}\n")
        p.add_run("Context: ").bold = True
        context = entry['context']
        match = re.search(r"\*\*(.+?)\*\*", context)
        if match:
            before, word, after = context.split("**", 2)
            r = p.add_run(before)
            r = p.add_run(word)
            r.font.highlight_color = 3  # Yellow highlight
            r = p.add_run(after)
        else:
            p.add_run(context)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def parse_banned_words(file):
    if file is None:
        return []
    content = file.read().decode("utf-8")
    return parse_banned_words_from_string(content)

def parse_banned_words_from_string(content):
    words = []
    for line in content.splitlines():
        line = line.strip()
        if line.startswith('"') and line.endswith('"'):
            words.append(line.strip('"'))
        elif line:
            words.append(line)
    return words

def load_default_banned_words_from_file(path="default_words.txt"):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return parse_banned_words_from_string(f.read())
    return []

# --- Precompiled Banned Words ---
DEFAULT_BANNED_WORDS = load_default_banned_words_from_file()

# --- Sidebar Inputs ---
st.sidebar.header("⚙️ Options")
use_default = st.sidebar.checkbox("Use default banned words", value=True)
text_input = st.sidebar.text_area("Enter additional banned words (one per line or quoted phrase)")
text_file = st.sidebar.file_uploader("Or upload a .txt file with banned words", type=["txt"])

# Optional: View default list
with st.sidebar.expander("📄 View default banned word list"):
    if DEFAULT_BANNED_WORDS:
        st.code("\n".join(DEFAULT_BANNED_WORDS))
    else:
        st.info("No default words loaded.")

# --- Main File Upload ---
uploaded_file = st.file_uploader("Upload a DOCX or PDF file", type=["docx", "pdf"])

if uploaded_file and (use_default or text_input or text_file):
    custom_words = [w.strip() for w in text_input.splitlines() if w.strip()]
    file_words = parse_banned_words(text_file)

    banned_words = DEFAULT_BANNED_WORDS if use_default else []
    banned_words += custom_words + file_words

    if uploaded_file.name.endswith(".docx"):
        text = extract_text_from_docx(uploaded_file)
    elif uploaded_file.name.endswith(".pdf"):
        uploaded_file.seek(0)
        text = extract_text_from_pdf(uploaded_file)
    else:
        st.error("Unsupported file type.")
        st.stop()

    st.subheader("📄 Full Text Preview")
    with st.expander("Show Document Text"):
        st.text_area("Extracted Text", value=text[:10000], height=300)

    st.subheader("🚫 Banned Words Found")
    results = highlight_banned_words(text, banned_words)

    if results:
        df = pd.DataFrame(results)
        st.dataframe(df, use_container_width=True)

        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button("Download Results as CSV", data=csv, file_name="banned_word_hits.csv")

        word_buffer = generate_word_doc(results, text, banned_words)
        st.download_button("Download Results as Word Document", data=word_buffer, file_name="banned_word_hits.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.success("No banned words found in the document.")

elif uploaded_file:
    st.info("Please enter or enable at least one banned word list to scan the document.")

# --- Footer Note ---
st.markdown("**Note:** Uploaded files and generated results are not stored on the server. All processing is done in-memory and data is only retained during your session.")
st.markdown("---")
st.markdown("If you find an issue or bug, please submit it [here](https://github.com/childrens-bti/word-scanner).")
